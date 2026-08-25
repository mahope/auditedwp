from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re, os

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

# --- Styles setup ---
style = doc.styles['Normal']
style.font.name = 'Georgia'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Title style
title_style = doc.styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
title_style.font.name = 'Georgia'
title_style.font.size = Pt(28)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x0B, 0x1A, 0x2A)
title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_style.paragraph_format.space_before = Pt(200)
title_style.paragraph_format.space_after = Pt(12)

subtitle_style = doc.styles.add_style('BookSubtitle', WD_STYLE_TYPE.PARAGRAPH)
subtitle_style.font.name = 'Georgia'
subtitle_style.font.size = Pt(16)
subtitle_style.font.color.rgb = RGBColor(0x4A, 0x5A, 0x6A)
subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_style.paragraph_format.space_after = Pt(200)

# Chapter styles
h1_style = doc.styles.add_style('Ch1', WD_STYLE_TYPE.PARAGRAPH)
h1_style.font.name = 'Georgia'
h1_style.font.size = Pt(22)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x0B, 0x1A, 0x2A)
h1_style.paragraph_format.space_before = Pt(36)
h1_style.paragraph_format.space_after = Pt(12)

h2_style = doc.styles.add_style('Ch2', WD_STYLE_TYPE.PARAGRAPH)
h2_style.font.name = 'Georgia'
h2_style.font.size = Pt(16)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x0B, 0x1A, 0x2A)
h2_style.paragraph_format.space_before = Pt(24)
h2_style.paragraph_format.space_after = Pt(8)

h3_style = doc.styles.add_style('Ch3', WD_STYLE_TYPE.PARAGRAPH)
h3_style.font.name = 'Georgia'
h3_style.font.size = Pt(13)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x4A, 0x5A, 0x6A)
h3_style.paragraph_format.space_before = Pt(18)
h3_style.paragraph_format.space_after = Pt(6)

with open('book/manuscript.md', 'r') as f:
    content = f.read()

lines = content.split('\n')
in_table = False
table_rows = []

for line in lines:
    stripped = line.strip()
    
    # Title page
    if stripped.startswith('# The Website Compliance Handbook'):
        doc.add_paragraph('The Website Compliance Handbook 2026', style='BookTitle')
        continue
    
    if stripped == '## GDPR, ePrivacy, NIS2 & Accessibility for European Small Business Owners':
        doc.add_paragraph(stripped.replace('## ', ''), style='BookSubtitle')
        continue
    
    if stripped.startswith('# '):
        continue  # top-level headings already handled
    
    if stripped.startswith('## '):
        doc.add_paragraph(stripped.replace('## ', ''), style='Ch1')
        continue
    
    if stripped.startswith('### '):
        doc.add_paragraph(stripped.replace('### ', ''), style='Ch2')
        continue
    
    if stripped.startswith('#### '):
        doc.add_paragraph(stripped.replace('#### ', ''), style='Ch3')
        continue
    
    if stripped == '---':
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        continue
    
    # Tables
    if stripped.startswith('|') and not stripped.startswith('|--'):
        in_table = True
        table_rows.append(stripped)
        continue
    else:
        if in_table and stripped.startswith('|--'):
            continue  # skip separator rows
        if in_table:
            # End of table — render
            if table_rows:
                headers = [h.strip() for h in table_rows[0].strip('|').split('|')]
                num_cols = len(headers)
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for i, row_text in enumerate(table_rows):
                    cells = [c.strip() for c in row_text.strip('|').split('|')]
                    for j, ctext in enumerate(cells):
                        if j < num_cols:
                            cell = table.rows[i].cells[j]
                            cell.text = ctext
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                doc.add_paragraph()
            in_table = False
            table_rows = []
    
    # List items
    if stripped.startswith('- **') and stripped.endswith('**'):
        doc.add_paragraph(stripped[3:].strip('*'), style='List Bullet')
        continue
    
    if stripped.startswith('- '):
        text = stripped[2:]
        if text.endswith('**'):
            text = text.strip('*')
        doc.add_paragraph(text, style='List Bullet')
        continue
    
    if re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        doc.add_paragraph(text, style='List Number')
        continue
    
    # Checklist items
    if stripped.startswith('- [ ] '):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run('\u2610 ' + stripped[6:])
        r.font.name = 'Segoe UI Symbol'
        continue
    if stripped.startswith('- [x] ') or stripped.startswith('- [X] '):
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run('\u2611 ' + stripped[6:])
        r.font.name = 'Segoe UI Symbol'
        continue
    
    # Regular paragraph with inline bold
    if stripped:
        parts = re.split(r'(\*\*.*?\*\*)', stripped)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part.strip('*'))
                run.font.bold = True
            else:
                p.add_run(part)

doc.save('book/The-Website-Compliance-Handbook-2026.docx')
print(f"DOCX saved. File: {os.path.getsize('book/The-Website-Compliance-Handbook-2026.docx')/1024:.1f} KB")
print(f"Word count: {len(content.split())}")
print(f"Est. pages: {len(content.split())//500 + 1}")