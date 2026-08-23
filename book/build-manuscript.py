#!/usr/bin/env python3
"""Convert the EU compliance guide from markdown to DOCX manuscript."""
import re, os

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    md_path = os.path.join(os.path.dirname(__file__), "eu-website-compliance-guide-2026.md")
    out_path = os.path.join(os.path.dirname(__file__), "eu-website-compliance-guide-2026.docx")

    with open(md_path) as f:
        content = f.read()

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, size in [(1, 22), (2, 16), (3, 12)]:
        hname = f'Heading {level}'
        if hname in doc.styles:
            h = doc.styles[hname]
            h.font.name = 'Times New Roman'
            h.font.size = Pt(size)
            h.font.bold = True
            h.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
            h.paragraph_format.space_after = Pt(8)
            h.paragraph_format.line_spacing = 1.2

    # ── Title page ──
    for _ in range(6):
        doc.add_paragraph('')

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tp.add_run('EU Website Compliance Guide 2026')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sp.add_run('The Practical Handbook for Website Owners\nServing EU Visitors')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x54, 0x67)
    run.font.italic = True

    doc.add_paragraph('')

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ap.add_run('By EUComply · August 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x47, 0x54, 0x67)

    doc.add_page_break()

    # ── Disclaimer ──
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dp.add_run('Disclaimer')
    run.font.size = Pt(12)
    run.font.bold = True

    dt = doc.add_paragraph()
    dt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dt.add_run(
        'This guide provides practical compliance information based on current EU regulations. '
        'It does not constitute legal advice. For high-risk compliance decisions, '
        'consult a qualified attorney licensed in your jurisdiction.'
    )
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ── TOC ──
    doc.add_heading('Table of Contents', level=1)
    toc = [
        ("1", "Introduction: Why Compliance Matters"),
        ("2", "GDPR Fundamentals for Websites"),
        ("3", "Cookie Consent — What the Law Actually Requires"),
        ("4", "Privacy Policy — The Art. 13 Checklist"),
        ("5", "Imprint / Legal Notice (Impressum)"),
        ("6", "Security Headers — Technical Compliance"),
        ("7", "Third-Party Trackers"),
        ("8", "European Accessibility Act (EAA) 2025"),
        ("9", "NIS2 and DORA — Resilience Requirements"),
        ("10", "Email Compliance — SPF, DKIM, DMARC"),
        ("11", "Compliance by Platform — Quick Reference"),
        ("12", "Maintaining Compliance — Monthly Checklist"),
        ("13", "Tools and Resources"),
        ("14", "Appendix: Regulation Overview Table"),
    ]
    for num, title in toc:
        p = doc.add_paragraph(f'Chapter {num}. {title}')
        p.paragraph_format.space_after = Pt(2)

    doc.add_page_break()

    # ── Body ──
    lines = content.split('\n')
    skip_toc = False
    in_chapter = False

    for line in lines:
        stripped = line.strip()

        # Skip YAML / front matter before first chapter
        if not in_chapter:
            if stripped.startswith('# 1.'):
                in_chapter = True
            else:
                continue

        # Skip TOC entries in markdown (we added our own)
        if stripped.startswith('- [') and '](#' in stripped:
            continue

        if stripped == '':
            continue

        # Headings
        if stripped.startswith('# ') and not stripped.startswith('# ' + '1'):
            txt = re.sub(r'^#+\s*', '', stripped)
            doc.add_heading(txt, level=1)
        elif stripped.startswith('## '):
            txt = re.sub(r'^##+\s*', '', stripped)
            doc.add_heading(txt, level=2)
        elif stripped.startswith('### '):
            txt = re.sub(r'^###+\s*', '', stripped)
            doc.add_heading(txt, level=3)
        elif stripped == '---':
            doc.add_paragraph('_' * 60)
        elif stripped.startswith('- [ ] '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(f'☐ {stripped[6:]}')
            run.font.size = Pt(10.5)
        elif stripped.startswith('| '):
            cells = [c.strip() for c in stripped.split('|') if c.strip()]
            if not all(c in ('---', ':', '---:') for c in cells):
                doc.add_paragraph(' | '.join(cells))
        else:
            # Remove bold markers
            plain = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
            # Remove checklist ** **
            if plain.startswith('**') and plain.endswith('**'):
                plain = plain.strip('*')
            if plain:
                doc.add_paragraph(plain)

    # ── Final page ──
    doc.add_page_break()
    for _ in range(4):
        doc.add_paragraph('')

    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run('— End of Guide —')
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x47, 0x54, 0x67)

    for url in [
        "Free compliance scan: https://auditedwp.pages.dev/scan/",
        "Automated daily monitoring: https://auditedwp.pages.dev/pro/"
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(url)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xD3)

    doc.save(out_path)
    wc = len(content.split())
    fs = os.path.getsize(out_path) // 1024
    print(f"✅ DOCX saved: {out_path}")
    print(f"   Words: {wc}  |  Est. pages: {max(30, wc // 350)}  |  File: {fs} KB")

if __name__ == '__main__':
    main()